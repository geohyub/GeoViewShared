#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) 2025 Junhyub Kim. All rights reserved.
"""
Design System -- Multi-Style Document Generation
==================================================
Central library for Word / PPT / Excel with 10 Word styles, 6 PPT styles,
4 Excel styles.  Each style is a self-contained palette + helper set.

Usage:
    from design_system import WORD_STYLES, build_word_document
    build_word_document(content, style="swiss_minimal", output="out.docx")

Style IDs:
    Word:  engineering, editorial, swiss_minimal, corporate,
           academic, dark_tech, warm_natural, korean_modern,
           geoview_report, fugro_report
    PPT:   dark_tech, clean_corporate, minimalist,
           bold_graphic, warm_natural, data_heavy
    Excel: clean_grid, dashboard, engineering_log, formal_report
"""

from __future__ import annotations
from dataclasses import dataclass, field
from typing import Dict, List, Tuple, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# 1 twip = 635 EMU (English Metric Unit)
_TWIP = 635


# ======================================================================
# Style Definition
# ======================================================================
@dataclass
class DocStyle:
    """Complete visual identity for a document."""
    id: str
    name: str
    name_kr: str
    concept: str
    concept_kr: str

    # Fonts
    font_heading: str = "Arial"
    font_body: str = "Arial"
    font_mono: str = "Consolas"
    font_size_h1: float = 18
    font_size_h2: float = 14
    font_size_h3: float = 0        # 0 = inherit from h2
    font_size_body: float = 10
    font_size_small: float = 8.5
    font_size_caption: float = 8

    # Colors (hex strings without #)
    color_primary: str = "000000"      # Main heading/accent
    color_h1: str = ""                 # H1-only color override (empty = use color_primary)
    color_caption: str = ""            # Caption color (empty = use color_body)
    color_secondary: str = "666666"    # Subheading/subtitle
    color_body: str = "333333"         # Body text
    color_accent: str = "E63946"       # Single highlight color
    color_bg_dark: str = "1A1D23"      # Dark background
    color_bg_light: str = "F5F5F5"     # Light background
    color_bg_alt: str = "FAFAFA"       # Alternating row
    color_border: str = "CCCCCC"       # Table/line borders
    color_table_header_bg: str = "333333"
    color_table_header_fg: str = "FFFFFF"

    # Spacing
    line_spacing: float = 1.35
    space_after_para: float = 8
    space_before_para: float = 0    # Body text space before (pt)
    space_before_h1: float = 24
    space_after_h1: float = 8
    body_indent_left: int = 0       # Body text left indent (twips)
    body_indent_right: int = 0      # Body text right indent (twips)
    body_first_line: int = 0        # Body text first-line indent (twips)
    body_char_spacing: int = 0      # Body text character spacing (half-points, e.g. -2)
    space_before_h2: float = 10     # Heading 2 space before (pt)
    space_after_h2: float = 10      # Heading 2 space after (pt)
    heading_justify: bool = False   # Headings use justified alignment
    heading_indent_left: int = 0    # H1 left indent (twips)
    heading_hanging: int = 0        # H1 hanging indent for numbers (twips)
    paper_a4: bool = False          # Force A4 paper size
    table_data_font_size: float = 0 # 0 = use font_size_body
    table_data_center: bool = False # Center-align data cells
    table_header_font_size: float = 0  # 0 = use font_size_body
    heading_indent_right: int = 0      # Heading right indent (twips)
    margin_top: float = 2.5     # cm
    margin_bottom: float = 2.5
    margin_left: float = 2.5
    margin_right: float = 2.5
    margin_header: float = 0    # cm, 0 = use default
    margin_footer: float = 0    # cm, 0 = use default

    # Footer
    page_footer_text: str = ""      # Footer left text (report title)
    page_footer_right: str = ""     # Footer right text (page numbering placeholder)

    # Style flags
    heading_italic: bool = False
    heading_uppercase: bool = False
    heading_border_bottom: bool = True
    heading_border_color: str = ""   # empty = use accent
    cover_style: str = "left"        # left, center, dark_block
    number_style: str = "bracket"    # bracket [01], dot 1., badge, none
    callout_style: str = "left_border"   # left_border, box, inline
    table_style: str = "minimal"     # minimal, filled, grid


# ======================================================================
# 10 Word Styles
# ======================================================================
WORD_STYLES: Dict[str, DocStyle] = {}

def _register(s: DocStyle):
    WORD_STYLES[s.id] = s
    return s

# --- 1. Engineering Blueprint ---
_register(DocStyle(
    id="engineering",
    name="Engineering Blueprint",
    name_kr="엔지니어링 블루프린트",
    concept="Technical drawing / spec sheet aesthetic",
    concept_kr="기술 도면 / 스펙시트 스타일. 모노스페이스 헤더, 다크 차콜 배경, 전기 시안 악센트.",
    font_heading="Consolas", font_body="Segoe UI", font_mono="Consolas",
    font_size_h1=16, font_size_h2=12, font_size_body=9.5,
    color_primary="00D4AA", color_secondary="888D96", color_body="444444",
    color_accent="00D4AA", color_bg_dark="1A1D23", color_bg_light="F7F8FA",
    color_bg_alt="F0F1F3", color_border="E2E5EA",
    color_table_header_bg="1A1D23", color_table_header_fg="00D4AA",
    line_spacing=1.3, margin_left=2.0, margin_right=2.0,
    heading_uppercase=True, heading_border_bottom=True, heading_border_color="00D4AA",
    cover_style="dark_block", number_style="bracket",
    callout_style="left_border", table_style="minimal",
))

# --- 2. Magazine Editorial ---
_register(DocStyle(
    id="editorial",
    name="Magazine Editorial",
    name_kr="매거진 에디토리얼",
    concept="Architecture journal / design magazine",
    concept_kr="건축 잡지 / 디자인 매거진 스타일. 세리프 이탤릭 헤더, 넓은 여백, 테라코타 악센트.",
    font_heading="Georgia", font_body="Segoe UI",
    font_size_h1=22, font_size_h2=15, font_size_body=10,
    color_primary="2C2420", color_secondary="9A8E82", color_body="443D36",
    color_accent="C4573A", color_bg_dark="3D3028", color_bg_light="FBF7F0",
    color_bg_alt="F5EEE4", color_border="E8E0D4",
    color_table_header_bg="FBF7F0", color_table_header_fg="C4573A",
    line_spacing=1.5, margin_top=3.0, margin_left=3.0, margin_right=3.0,
    heading_italic=True, heading_border_bottom=False,
    cover_style="center", number_style="dot",
    callout_style="left_border", table_style="minimal",
))

# --- 3. Swiss Minimal ---
_register(DocStyle(
    id="swiss_minimal",
    name="Swiss Minimal",
    name_kr="스위스 미니멀",
    concept="Bauhaus / Swiss railway timetable",
    concept_kr="바우하우스 / 스위스 그리드 디자인. 블랙+화이트+레드 3색만, 엄격한 그리드.",
    font_heading="Arial", font_body="Arial",
    font_size_h1=18, font_size_h2=13, font_size_body=10,
    color_primary="000000", color_secondary="6B6B6B", color_body="333333",
    color_accent="E63946", color_bg_dark="000000", color_bg_light="F5F5F5",
    color_bg_alt="F5F5F5", color_border="CCCCCC",
    color_table_header_bg="000000", color_table_header_fg="FFFFFF",
    line_spacing=1.45, margin_left=3.5, margin_right=2.5,
    heading_uppercase=False, heading_border_bottom=True, heading_border_color="000000",
    cover_style="left", number_style="badge",
    callout_style="box", table_style="minimal",
))

# --- 4. Corporate Standard ---
_register(DocStyle(
    id="corporate",
    name="Corporate Standard",
    name_kr="기업 표준",
    concept="Professional navy-blue business document",
    concept_kr="네이비 블루 기업 문서. 가장 무난하고 범용적인 비즈니스 스타일.",
    font_heading="Calibri", font_body="Calibri",
    font_size_h1=18, font_size_h2=14, font_size_body=11,
    color_primary="1B3A5C", color_secondary="4A6D8C", color_body="333333",
    color_accent="2A82DA", color_bg_dark="1B3A5C", color_bg_light="F0F4F8",
    color_bg_alt="E8EDF3", color_border="BCC8D6",
    color_table_header_bg="1B3A5C", color_table_header_fg="FFFFFF",
    line_spacing=1.15, margin_top=2.5, margin_left=2.5, margin_right=2.5,
    heading_border_bottom=True, heading_border_color="2A82DA",
    cover_style="center", number_style="dot",
    callout_style="left_border", table_style="filled",
))

# --- 5. Academic Research ---
_register(DocStyle(
    id="academic",
    name="Academic Research",
    name_kr="학술 논문",
    concept="Formal research paper / thesis style",
    concept_kr="학술 논문 / 연구 보고서 스타일. Times 세리프, 좁은 줄간격, 형식적.",
    font_heading="Times New Roman", font_body="Times New Roman",
    font_size_h1=14, font_size_h2=12, font_size_body=11,
    color_primary="000000", color_secondary="444444", color_body="000000",
    color_accent="8B0000", color_bg_dark="333333", color_bg_light="F9F9F9",
    color_bg_alt="F0F0F0", color_border="999999",
    color_table_header_bg="F0F0F0", color_table_header_fg="000000",
    line_spacing=1.15, space_after_para=6,
    margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=2.54,
    heading_border_bottom=False,
    cover_style="center", number_style="dot",
    callout_style="inline", table_style="grid",
))

# --- 6. Dark Tech Report ---
_register(DocStyle(
    id="dark_tech",
    name="Dark Tech Report",
    name_kr="다크 테크 리포트",
    concept="Dark IDE / terminal-inspired report",
    concept_kr="IDE / 터미널 영감의 다크 리포트. 코드 블록 느낌, 네온 그린 악센트.",
    font_heading="Cascadia Code", font_body="Segoe UI", font_mono="Cascadia Code",
    font_size_h1=16, font_size_h2=13, font_size_body=9.5,
    color_primary="4EC9B0", color_secondary="7A7E85", color_body="D4D4D4",
    color_accent="DCDCAA", color_bg_dark="1E1E1E", color_bg_light="2D2D2D",
    color_bg_alt="252526", color_border="404040",
    color_table_header_bg="1E1E1E", color_table_header_fg="4EC9B0",
    line_spacing=1.35,
    margin_top=2.0, margin_left=2.0, margin_right=2.0,
    heading_uppercase=True, heading_border_bottom=True, heading_border_color="4EC9B0",
    cover_style="dark_block", number_style="bracket",
    callout_style="box", table_style="minimal",
))

# --- 7. Warm Natural ---
_register(DocStyle(
    id="warm_natural",
    name="Warm Natural",
    name_kr="따뜻한 자연",
    concept="Earth tones / organic / hand-crafted feel",
    concept_kr="어스톤 / 유기적 디자인. 올리브 그린과 머스타드, 자연스러운 느낌.",
    font_heading="Cambria", font_body="Segoe UI",
    font_size_h1=20, font_size_h2=14, font_size_body=10,
    color_primary="4A5D23", color_secondary="8B7D5E", color_body="3D3B30",
    color_accent="C17817", color_bg_dark="3D3B30", color_bg_light="F8F5EE",
    color_bg_alt="F0EBDF", color_border="D4CBBA",
    color_table_header_bg="4A5D23", color_table_header_fg="FFFFFF",
    line_spacing=1.45,
    margin_top=2.5, margin_left=2.8, margin_right=2.8,
    heading_italic=False, heading_border_bottom=True, heading_border_color="C17817",
    cover_style="center", number_style="dot",
    callout_style="left_border", table_style="filled",
))

# --- 8. Korean Modern ---
_register(DocStyle(
    id="korean_modern",
    name="Korean Modern",
    name_kr="한국 모던",
    concept="Korean design aesthetic / clean geometry / hangeul-optimized",
    concept_kr="한글 최적화 모던 디자인. 맑은 고딕, 깔끔한 기하학, 짙은 잉크 블루 악센트.",
    font_heading="Malgun Gothic", font_body="Malgun Gothic",
    font_size_h1=17, font_size_h2=13, font_size_body=10,
    color_primary="1A2744", color_secondary="5C6B80", color_body="2D2D2D",
    color_accent="D4483B", color_bg_dark="1A2744", color_bg_light="F4F6F9",
    color_bg_alt="EDF0F5", color_border="C5CDD8",
    color_table_header_bg="1A2744", color_table_header_fg="FFFFFF",
    line_spacing=1.6,   # Korean text needs more line spacing
    space_after_para=8,
    margin_top=2.5, margin_left=2.5, margin_right=2.5,
    heading_border_bottom=True, heading_border_color="D4483B",
    cover_style="left", number_style="dot",
    callout_style="left_border", table_style="filled",
))

# --- 9. GeoView Client Report ---
_register(DocStyle(
    id="geoview_report",
    name="GeoView Client Report",
    name_kr="지오뷰 클라이언트 보고서",
    concept="Formal client-facing geophysical survey processing report",
    concept_kr="해양 물리탐사 보고서. Arial 기반, 스틸블루 헤더, 풀 그리드 테이블, 정식 리비전 관리 포함.",
    font_heading="Arial", font_body="Arial",
    font_size_h1=14, font_size_h2=14, font_size_h3=13, font_size_body=11,
    font_size_small=9, font_size_caption=10,
    color_primary="000000",       # Black headings
    color_secondary="4F81BD",     # Steel blue for project name / accent
    color_body="000000",          # Black body text
    color_accent="4F81BD",        # Steel blue accent
    color_bg_dark="4E81BD",       # Table header fill (exact from real docs)
    color_bg_light="FFFFFF",      # White background
    color_bg_alt="F2F7FB",        # Subtle blue-tint alternating rows
    color_border="000000",        # Black grid borders
    color_table_header_bg="4E81BD",  # Exact match: real docs use 4E81BD
    color_table_header_fg="FFFFFF",
    line_spacing=1.0,             # Single spacing (real docs use single, not 1.15)
    space_after_para=12,          # 240 twips = 12pt
    space_before_para=8.15,       # 163 twips = 8.15pt (Body Text style exact)
    space_before_h1=12,           # 240 twips = 12pt
    space_after_h1=10,            # 200 twips = 10pt
    space_before_h2=10,           # 200 twips = 10pt
    space_after_h2=10,            # 200 twips = 10pt
    body_indent_left=220,         # 220 twips ≈ 3.9mm each side
    body_indent_right=220,
    body_first_line=57,           # 57 twips ≈ 1mm first-line indent
    body_char_spacing=-2,         # Tight character spacing
    heading_justify=True,         # All headings use justified alignment
    heading_indent_left=284,      # H1 hanging indent base
    heading_hanging=284,          # H1 hanging indent for number
    heading_indent_right=220,     # All headings right indent = 220tw (rightChars=100)
    paper_a4=True,                # A4 paper size
    table_data_font_size=10,      # 10pt for table data cells
    table_data_center=True,       # Center-align all data cells
    table_header_font_size=10,    # 10pt Bold for table header cells (real docs: sz=20)
    margin_top=2.89,              # 1639 twips ≈ 2.89 cm
    margin_bottom=2.75,           # ~2.75 cm
    margin_left=2.3,              # 1304 twips ≈ 2.3 cm
    margin_right=2.3,
    margin_header=1.5,            # 851 twips = 1.50 cm (real docs exact)
    margin_footer=1.5,            # 850 twips = 1.50 cm
    heading_italic=False,
    heading_uppercase=True,        # H1 is ALL CAPS in real GeoView docs
    heading_border_bottom=False,   # No border -- clean heading style
    cover_style="geoview",        # Special GeoView cover
    number_style="dot",           # "1. Introduction" style
    callout_style="inline",       # No decorative callouts
    table_style="grid",           # Full grid borders
))

# --- 10. Fugro Report ---
_register(DocStyle(
    id="fugro_report",
    name="Fugro Technical Report",
    name_kr="Fugro 기술 보고서",
    concept="Fugro offshore survey report template — Segoe UI, dark navy + slate blue",
    concept_kr="Fugro 해양조사 보고서. Segoe UI, 다크 네이비 브랜드 컬러, 슬레이트 블루 H1, 그레이 캡션.",
    font_heading="Segoe UI", font_body="Segoe UI", font_mono="Consolas",
    font_size_h1=18, font_size_h2=13, font_size_h3=11, font_size_body=11,
    font_size_small=8, font_size_caption=9,
    color_primary="000000",       # H2/H3/body: black
    color_h1="6788B1",            # H1: slate blue
    color_caption="7F7F7F",       # Captions: medium gray
    color_secondary="011E41",     # Dark navy (brand, header/footer)
    color_body="000000",
    color_accent="6788B1",        # Slate blue accent
    color_bg_dark="011E41",       # Fugro dark navy
    color_bg_light="FFFFFF",
    color_bg_alt="F5F5F5",        # Light alternating row (subtle)
    color_border="CCCCCC",        # Light borders
    color_table_header_bg="6788B1",   # Slate blue header fill
    color_table_header_fg="FFFFFF",   # White header text
    line_spacing=1.15,            # Slightly open (not single, not 1.5)
    space_after_para=8,
    space_before_para=0,
    space_before_h1=18,
    space_after_h1=6,
    space_before_h2=12,
    space_after_h2=4,
    body_indent_left=0,           # No body indent (clean style)
    body_indent_right=0,
    body_first_line=0,
    body_char_spacing=0,
    heading_justify=False,        # Left-aligned headings
    heading_indent_left=0,
    heading_hanging=0,
    heading_indent_right=0,
    paper_a4=True,
    table_data_font_size=9,       # Compact 9pt tables
    table_data_center=False,      # Left-aligned data (row labels bold)
    table_header_font_size=9,     # 9pt header text
    margin_top=2.5,
    margin_bottom=2.5,
    margin_left=2.0,
    margin_right=2.0,
    margin_header=1.0,
    margin_footer=1.0,
    heading_italic=False,
    heading_uppercase=False,       # Fugro does NOT use ALL CAPS
    heading_border_bottom=False,   # No heading border
    cover_style="center",         # Clean centered cover
    number_style="dot",           # "1. Section" style
    callout_style="left_border",
    table_style="filled",         # Filled header + subtle rows
))


# ======================================================================
# XML Helpers (shared)
# ======================================================================
def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shading(cell, fill_hex):
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    ))


def _valign(cell, val="center"):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    )


def _para_border(para, side, sz, color, space=4):
    pPr = para._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:{side} w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    ))


def _para_shading(para, fill_hex):
    pPr = para._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
    ))


def _run_shading(run, fill_hex):
    rPr = run._element.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
    ))


def _set_full_font_on_style(style, font_name):
    """Set font name on all four font slots (ascii, hAnsi, eastAsia, cs) for a style."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), font_name)


def _cell_border(cell, sides, sz=4, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    xml = f'<w:tcBorders {nsdecls("w")}>'
    for side in sides:
        xml += f'<w:{side} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
    xml += '</w:tcBorders>'
    tcPr.append(parse_xml(xml))


# ======================================================================
# Word Document Builder (style-aware)
# ======================================================================
class WordBuilder:
    """Build a Word document using a specific DocStyle."""

    def __init__(self, style: DocStyle):
        self.s = style
        self.doc = Document()
        self._fig = 0
        self._tbl = 0   # Table counter for captions

        # Setup named styles BEFORE any content is added
        self._setup_heading_styles()
        self._setup_body_style()
        self._setup_caption_style()

        # Page setup
        for sec in self.doc.sections:
            sec.top_margin = Cm(style.margin_top)
            sec.bottom_margin = Cm(style.margin_bottom)
            sec.left_margin = Cm(style.margin_left)
            sec.right_margin = Cm(style.margin_right)
            # A4 paper size (11906 x 16838 twips)
            if style.paper_a4:
                sec.page_width = Cm(21.0)   # 11906 twips
                sec.page_height = Cm(29.7)  # 16838 twips
            # Header / footer margins
            if style.margin_header:
                sec.header_distance = Cm(style.margin_header)
            if style.margin_footer:
                sec.footer_distance = Cm(style.margin_footer)

    # --- Style Definition Setup ---
    def _setup_heading_styles(self):
        """Customize built-in Heading 1/2/3 style definitions to match DocStyle.

        This ensures TOC, Navigation Pane, and Outline View reflect correct
        formatting. Only numbering prefix and hanging indent need direct
        formatting when adding headings.
        """
        s = self.s
        style_configs = {
            1: {
                'size': s.font_size_h1,
                'space_before': s.space_before_h1,
                'space_after': s.space_after_h1,
                'all_caps': s.heading_uppercase,
                'italic': s.heading_italic,
            },
            2: {
                'size': s.font_size_h2,
                'space_before': s.space_before_h2,
                'space_after': s.space_after_h2,
                'all_caps': False,
                'italic': False,
            },
            3: {
                'size': s.font_size_h3 if s.font_size_h3 else s.font_size_h2,
                'space_before': s.space_before_h2,
                'space_after': s.space_after_h2,
                'all_caps': False,
                'italic': False,
            },
        }
        for level, cfg in style_configs.items():
            style = self.doc.styles[f'Heading {level}']
            # Character formatting
            font = style.font
            font.name = s.font_heading
            font.size = Pt(cfg['size'])
            font.bold = True
            font.italic = cfg['italic']
            # H1 can have a different color (e.g., Fugro slate blue)
            h_color = (s.color_h1 if s.color_h1 and level == 1 else s.color_primary)
            font.color.rgb = _rgb(h_color)
            font.all_caps = cfg['all_caps']
            _set_full_font_on_style(style, s.font_heading)
            # Paragraph formatting
            pf = style.paragraph_format
            pf.space_before = Pt(cfg['space_before'])
            pf.space_after = Pt(cfg['space_after'])
            pf.line_spacing = s.line_spacing
            pf.keep_with_next = True
            if s.heading_justify:
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _setup_body_style(self):
        """Create or customize the 'Body Text' paragraph style.

        Defines justified body text with correct font, spacing, indentation,
        and character spacing. Using a named style means body_text() needs
        no direct formatting overrides.
        """
        s = self.s
        try:
            style = self.doc.styles['Body Text']
        except KeyError:
            style = self.doc.styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = self.doc.styles['Normal']
        # Character formatting
        font = style.font
        font.name = s.font_body
        font.size = Pt(s.font_size_body)
        font.color.rgb = _rgb(s.color_body)
        font.bold = False
        font.italic = False
        _set_full_font_on_style(style, s.font_body)
        # Character spacing (not exposed by python-docx API)
        if s.body_char_spacing:
            rPr = style.element.get_or_add_rPr()
            existing = rPr.find(qn('w:spacing'))
            if existing is not None:
                rPr.remove(existing)
            rPr.append(parse_xml(
                f'<w:spacing {nsdecls("w")} w:val="{s.body_char_spacing}"/>'
            ))
        # Paragraph formatting
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(s.space_before_para)
        pf.space_after = Pt(s.space_after_para)
        pf.line_spacing = s.line_spacing
        if s.body_indent_left:
            pf.left_indent = s.body_indent_left * _TWIP
        if s.body_indent_right:
            pf.right_indent = s.body_indent_right * _TWIP
        if s.body_first_line:
            pf.first_line_indent = s.body_first_line * _TWIP
        # Set heading → body text transition
        for level in (1, 2, 3):
            self.doc.styles[f'Heading {level}'].next_paragraph_style = style
        self._body_style = style

    def _setup_caption_style(self):
        """Create 'GV Caption' style for table/figure captions."""
        s = self.s
        style_name = 'GV Caption'
        try:
            style = self.doc.styles[style_name]
        except KeyError:
            style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = self.doc.styles['Normal']
        font = style.font
        font.name = s.font_body
        font.size = Pt(s.font_size_caption)
        font.bold = True
        cap_color = s.color_caption if s.color_caption else s.color_body
        font.color.rgb = _rgb(cap_color)
        _set_full_font_on_style(style, s.font_body)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing = s.line_spacing
        self._caption_style = style

    def _run(self, para, text, font=None, size=None, color=None,
             bold=False, italic=False):
        """Add a formatted run."""
        run = para.add_run(text)
        run.font.name = font or self.s.font_body
        run.font.size = Pt(size or self.s.font_size_body)
        run.font.color.rgb = _rgb(color or self.s.color_body)
        run.bold = bold
        run.italic = italic
        return run

    # --- Cover ---
    def cover(self, title, subtitle, meta=""):
        s = self.s
        if s.cover_style == "dark_block":
            self._cover_dark(title, subtitle, meta)
        elif s.cover_style == "geoview":
            self._cover_geoview(title, subtitle, meta)
        elif s.cover_style == "center":
            self._cover_center(title, subtitle, meta)
        else:  # "left"
            self._cover_left(title, subtitle, meta)

    def _cover_dark(self, title, subtitle, meta):
        s = self.s
        # Dark background block
        for text, sz, color in [
            ("\n", 20, s.color_secondary),
            (f"  {title.upper()}" if s.heading_uppercase else f"  {title}", 28, s.color_accent),
            (f"  {subtitle}", 13, s.color_secondary),
            (f"  {meta}", 9, s.color_secondary),
            (" ", 8, s.color_secondary),
        ]:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _para_shading(p, s.color_bg_dark)
            run = self._run(p, text, font=s.font_heading, size=sz, color=color,
                           bold=(sz > 20))

        # Accent line
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(20)
        _para_border(p, "top", 10, s.color_accent, space=0)
        p.add_run(" ").font.size = Pt(1)

    def _cover_center(self, title, subtitle, meta):
        s = self.s
        # Whitespace
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        self._run(p, title, font=s.font_heading, size=32,
                  color=s.color_primary, bold=True, italic=s.heading_italic)

        # Thin accent line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _para_border(p, "bottom", 4, s.color_accent, space=0)
        p.add_run(" " * 30).font.size = Pt(2)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self._run(p, subtitle, size=11, color=s.color_secondary)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        self._run(p, meta, size=9, color=s.color_secondary)

        self.doc.add_page_break()

    def _cover_left(self, title, subtitle, meta):
        s = self.s
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(16)

        # Accent bar
        p = self.doc.add_paragraph()
        _para_shading(p, s.color_accent)
        p.paragraph_format.space_after = Pt(0)
        p.add_run("  ").font.size = Pt(4)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        self._run(p, title, font=s.font_heading, size=36,
                  color=s.color_primary, bold=True)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        self._run(p, subtitle.upper() if s.heading_uppercase else subtitle,
                  size=12, color=s.color_secondary)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(40)
        self._run(p, meta, size=9, color=s.color_secondary)

        self.doc.add_page_break()

    def _cover_geoview(self, title, subtitle, meta):
        """GeoView-style centered cover with steel blue project name and black revision."""
        s = self.s
        # Spacing before title
        for _ in range(6):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(14)

        # Title (black, 24pt bold centered)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        self._run(p, title, font="Arial", size=24, color="000000", bold=True)

        # Project name lines (steel blue, 24pt bold centered)
        for line in subtitle.split("\n"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            self._run(p, line, font="Arial", size=24, color=s.color_secondary, bold=True)

        # Spacing
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(12)

        # Revision / Date (black, 18pt bold centered)
        if meta:
            for line in meta.split("\n"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                self._run(p, line, font="Arial", size=18, color="000000", bold=True)

        # More spacing
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(10)

        # Company info block (small, centered)
        company_lines = [
            "GEOVIEW CO., Ltd",
            "2, 4, 5, 7th Floor, Sejin Building,",
            "423 Hasinbeonyeong-Ro, Saha-gu,",
            "Busan City, Republic of Korea",
        ]
        for line in company_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            self._run(p, line, font="Arial", size=10, color="000000",
                     bold=(line == company_lines[0]))

        self.doc.add_page_break()

    # --- Front matter (GeoView-specific) ---
    def revision_history(self, revisions):
        """Add Revision History table.
        revisions: list of (rev, date, status, check, approval) tuples.
        """
        headers = ["REVISION", "DATE", "STATUS", "CHECK", "APPROVAL"]
        self._geoview_label("REVISION HISTORY")
        self.table(headers, revisions)
        return

    def revision_log(self, entries):
        """Add Revision Log table.
        entries: list of (date, section, change) tuples.
        """
        headers = ["DATE", "SECTION", "CHANGE"]
        self._geoview_label("REVISION LOG")
        self.table(headers, entries)
        return

    def document_control(self, roles):
        """Add Document Control table.
        roles: list of (responsibility, position, name) tuples.
        """
        headers = ["RESPONSIBILITY", "POSITION", "NAME"]
        self._geoview_label("DOCUMENT CONTROL")
        self.table(headers, roles)
        return

    def abbreviations(self, abbr_pairs):
        """Add abbreviations table.
        abbr_pairs: list of (abbreviation, description) tuples.
        """
        headers = ["Abbreviation", "Description"]
        self.heading("Abbreviations", level=2)
        t = self.table(headers, abbr_pairs)
        self.table_caption("Table 1. Abbreviations")
        return t

    def figure_caption(self, text, number=None):
        """Add a figure caption (centered, in a single-cell borderless table)."""
        self._fig += 1
        num = number or self._fig
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = self.doc.styles['GV Caption']
        p.add_run(f"Figure {num}. {text}")
        _cell_margins(cell, top=30, bottom=30, left=60, right=60)
        for side in ["top", "bottom", "left", "right"]:
            _cell_border(cell, [side], sz=0, color="FFFFFF")
        return p

    def table_caption(self, text, number=None):
        """Add a table caption below a table (centered, in a single-cell borderless table)."""
        self._tbl += 1
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = self.doc.styles['GV Caption']
        p.add_run(text)
        _cell_margins(cell, top=30, bottom=30, left=60, right=60)
        for side in ["top", "bottom", "left", "right"]:
            _cell_border(cell, [side], sz=0, color="FFFFFF")
        return p

    def _geoview_label(self, text):
        """Centered section label (for front-matter blocks)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        self._run(p, text, font="Arial", size=12, color="000000", bold=True)
        return p

    def body_text(self, text):
        """Add justified body text using the 'Body Text' named style.

        All formatting (font, size, color, spacing, indentation, character
        spacing) is defined in the style. No direct overrides needed.
        """
        p = self.doc.add_paragraph(style='Body Text')
        p.add_run(text)
        return p

    # --- Section Heading ---
    def heading(self, text, level=1, number=None):
        """Add a section heading with optional hierarchical numbering.

        Uses Word's built-in Heading 1/2/3 styles (pre-configured in
        _setup_heading_styles) for proper TOC, Navigation Pane, and Outline
        support. Only numbering prefix and hanging indent are applied as
        direct formatting.
        """
        s = self.s

        # --- Auto-numbering logic ---
        auto_prefix = ""
        if hasattr(self, '_h_counters') and number is not False:
            if level == 1:
                self._h_counters[0] += 1
                self._h_counters[1] = 0
                self._h_counters[2] = 0
                auto_prefix = f"{self._h_counters[0]}. "
            elif level == 2:
                self._h_counters[1] += 1
                self._h_counters[2] = 0
                auto_prefix = f"{self._h_counters[0]}.{self._h_counters[1]} "
            elif level == 3:
                self._h_counters[2] += 1
                auto_prefix = (
                    f"{self._h_counters[0]}.{self._h_counters[1]}"
                    f".{self._h_counters[2]} "
                )

        # Create paragraph with built-in heading style (for TOC)
        style_name = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}.get(level)
        p = self.doc.add_paragraph(style=style_name) if style_name else self.doc.add_paragraph()

        # Hanging indent (direct formatting — varies by level)
        r_attr = f' w:right="{s.heading_indent_right}"' if s.heading_indent_right else ''
        if s.heading_indent_left and level == 1:
            pPr = p._element.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:ind {nsdecls("w")} w:left="{s.heading_indent_left}" '
                f'w:hanging="{s.heading_hanging}"{r_attr}/>'
            ))
        elif s.heading_indent_left and level == 2:
            pPr = p._element.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:ind {nsdecls("w")} w:left="{s.heading_indent_left + 284}" '
                f'w:hanging="{s.heading_hanging + 142}"{r_attr}/>'
            ))
        elif s.heading_indent_left and level == 3:
            pPr = p._element.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:ind {nsdecls("w")} w:left="{s.heading_indent_left + 568}" '
                f'w:hanging="{s.heading_hanging + 284}"{r_attr}/>'
            ))

        # --- Add text content ---
        if auto_prefix:
            # Auto-numbered: single run, inherits all formatting from style
            # (caps, font, size, color, bold are all in the style definition)
            p.add_run(auto_prefix + text)
        elif number is not None and number is not False and not hasattr(self, '_h_counters'):
            # Legacy manual numbering (non-auto styles)
            ns = s.number_style
            sz = s.font_size_h1 if level == 1 else s.font_size_h2
            if ns == "bracket":
                prefix = f"[{number:02d}]  "
            elif ns == "badge":
                badge = self._run(p, f" {number} ", font=s.font_heading,
                                  size=sz * 0.65, color=s.color_table_header_fg, bold=True)
                _run_shading(badge, s.color_accent)
                self._run(p, "  ", size=sz)
                prefix = None
            else:
                prefix = f"{number}.  "
            if prefix:
                self._run(p, prefix, font=s.font_heading,
                          size=sz * 0.7, color=s.color_accent, bold=True)
            display_text = text.upper() if s.heading_uppercase and level == 1 else text
            self._run(p, display_text, font=s.font_heading, size=sz,
                      color=s.color_primary, bold=True,
                      italic=s.heading_italic and level == 1)
        else:
            # No numbering: plain run, inherits everything from style
            p.add_run(text)

        # Bottom border
        bc = s.heading_border_color or s.color_accent
        if s.heading_border_bottom and level == 1:
            _para_border(p, "bottom", 4, bc, space=3)

        return p

    def reset_numbering(self):
        """Enable auto-numbering for headings. Call this once before the first heading."""
        self._h_counters = [0, 0, 0]

    # --- Body paragraph ---
    def para(self, text, bold=False, italic=False):
        s = self.s
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(s.space_after_para)
        p.paragraph_format.line_spacing = s.line_spacing
        self._run(p, text, bold=bold, italic=italic)
        return p

    # --- Callout / Note ---
    def callout(self, text, kind="NOTE"):
        s = self.s
        labels = {"NOTE": "Note", "TIP": "Tip", "WARNING": "Warning"}
        colors_map = {
            "NOTE":    (s.color_accent,  "EBF3FC"),
            "TIP":     ("28A745",        "E8F5E9"),
            "WARNING": ("E67E22",        "FFF3E0"),
        }
        border_hex, bg_hex = colors_map.get(kind, colors_map["NOTE"])

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = s.line_spacing
        p.paragraph_format.left_indent = Inches(0.15)

        label = self._run(p, f"  {labels[kind]}:  ", font=s.font_body,
                          size=s.font_size_small, color=border_hex, bold=True)
        self._run(p, text, size=s.font_size_small)

        if s.callout_style == "left_border":
            _para_border(p, "left", 20, border_hex, space=8)
        elif s.callout_style == "box":
            _para_shading(p, bg_hex)
            _para_border(p, "left", 20, border_hex, space=8)
        # "inline" = no special decoration

        return p

    # --- Table ---
    def table(self, headers, rows):
        s = self.s
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # --- Table-level borders for "grid" and "filled" styles ---
        if s.table_style in ("grid", "filled"):
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            # Full grid borders (all 6 directions)
            borders_xml = (
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{s.color_border}"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(parse_xml(borders_xml))
            # Table width = 100%
            tblPr.append(parse_xml(
                f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
            ))

        # Header font size
        h_font_sz = s.table_header_font_size if s.table_header_font_size else s.font_size_small
        # Data font size
        d_font_sz = s.table_data_font_size if s.table_data_font_size else (s.font_size_body * 0.9)

        # Header row
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if s.table_style != "minimal" else WD_ALIGN_PARAGRAPH.LEFT
            is_filled = s.table_style in ("filled", "grid")

            if is_filled:
                _shading(cell, s.color_table_header_bg)
                fg = s.color_table_header_fg
            else:
                fg = s.color_accent

            run = p.add_run(h.upper() if s.heading_uppercase else h)
            run.font.name = s.font_mono if s.table_style == "minimal" and "Consolas" in s.font_mono else s.font_body
            run.font.size = Pt(h_font_sz)
            run.font.color.rgb = _rgb(fg)
            run.bold = True
            if s.table_style in ("grid", "filled"):
                _cell_margins(cell, top=45, bottom=45, left=115, right=115)
            else:
                _cell_margins(cell, top=35, bottom=35, left=70, right=70)
            _valign(cell, "center")

            if not is_filled:
                _cell_border(cell, ["bottom"], sz=6, color=s.color_accent)

        # Data rows
        for r, row_data in enumerate(rows):
            for j, val in enumerate(row_data):
                cell = table.rows[r + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                if s.table_data_center:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                run.font.name = s.font_body
                run.font.size = Pt(d_font_sz)
                run.font.color.rgb = _rgb(s.color_body)
                if j == 0 and not s.table_data_center:
                    run.bold = True

                if s.table_style in ("grid", "filled"):
                    _cell_margins(cell, top=35, bottom=35, left=115, right=115)
                else:
                    _cell_margins(cell, top=25, bottom=25, left=70, right=70)
                _valign(cell, "center")

                if s.table_style == "filled" and r % 2 == 0:
                    _shading(cell, s.color_bg_alt)
                elif s.table_style == "grid":
                    if r % 2 == 0 and s.color_bg_alt.upper() != "FFFFFF":
                        _shading(cell, s.color_bg_alt)
                else:  # minimal
                    _cell_border(cell, ["bottom"], sz=1, color=s.color_border)

        return table

    # --- Numbered list ---
    def numbered_list(self, items):
        s = self.s
        for idx, item in enumerate(items):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.line_spacing = s.line_spacing

            if s.number_style == "bracket":
                num_text = f" {idx+1} "
                num_run = self._run(p, num_text, font=s.font_mono,
                                    size=s.font_size_small, color="FFFFFF", bold=True)
                _run_shading(num_run, s.color_bg_dark)
                self._run(p, "  ", size=s.font_size_body)
            elif s.number_style == "badge":
                num_run = self._run(p, f" {idx+1} ", font=s.font_heading,
                                    size=s.font_size_small, color="FFFFFF", bold=True)
                _run_shading(num_run, s.color_accent)
                self._run(p, "  ", size=s.font_size_body)
            else:  # dot
                num_run = self._run(p, f"{idx+1}.  ", font=s.font_heading,
                                    size=s.font_size_h2, color=s.color_accent,
                                    bold=True, italic=s.heading_italic)

            self._run(p, item)
        return

    # --- Bullet list ---
    def bullet_list(self, items):
        s = self.s
        bullets = {"engineering": ">", "editorial": "--", "swiss_minimal": "\u25A0",
                   "dark_tech": "$", "academic": "\u2022", "korean_modern": "\u25B6"}
        bullet_char = bullets.get(s.id, "\u2022")

        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.line_spacing = s.line_spacing

            b = self._run(p, f"{bullet_char}  ", size=s.font_size_small,
                          color=s.color_accent, bold=True)
            self._run(p, item)
        return

    # --- Key-Value pairs (like shortcuts) ---
    def kv_table(self, pairs):
        """Two-column key-value table (keys right-aligned, colored)."""
        s = self.s
        table = self.doc.add_table(rows=len(pairs), cols=2)
        for r, (key, val) in enumerate(pairs):
            cell_k = table.rows[r].cells[0]
            cell_k.text = ""
            p = cell_k.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(key)
            run.font.name = s.font_mono
            run.font.size = Pt(s.font_size_body * 0.9)
            run.font.color.rgb = _rgb(s.color_accent)
            run.bold = True
            if s.id in ("engineering", "dark_tech"):
                _run_shading(run, s.color_bg_dark)
            _cell_margins(cell_k, top=20, bottom=20, left=60, right=15)
            _cell_border(cell_k, ["bottom"], sz=1, color=s.color_border)

            cell_v = table.rows[r].cells[1]
            cell_v.text = ""
            p = cell_v.paragraphs[0]
            run = p.add_run(val)
            run.font.name = s.font_body
            run.font.size = Pt(s.font_size_body * 0.9)
            run.font.color.rgb = _rgb(s.color_body)
            _cell_margins(cell_v, top=20, bottom=20, left=15, right=60)
            _cell_border(cell_v, ["bottom"], sz=1, color=s.color_border)

        return table

    # --- Horizontal rule ---
    def hr(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _para_border(p, "bottom", 4, self.s.color_accent, space=1)
        p.add_run(" ").font.size = Pt(1)
        return p

    # --- Footer ---
    def footer(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(30)
        _para_border(p, "top", 2, self.s.color_border, space=6)
        self._run(p, text, size=7.5, color=self.s.color_secondary,
                  italic=True)
        return p

    # --- Page break ---
    def page_break(self):
        self.doc.add_page_break()

    # --- Section break (for separate header/footer per section) ---
    def section_break(self):
        """Add a section break (new page) to allow different headers/footers."""
        from docx.enum.section import WD_ORIENT
        new_sec = self.doc.add_section()
        s = self.s
        new_sec.top_margin = Cm(s.margin_top)
        new_sec.bottom_margin = Cm(s.margin_bottom)
        new_sec.left_margin = Cm(s.margin_left)
        new_sec.right_margin = Cm(s.margin_right)
        if s.paper_a4:
            new_sec.page_width = Cm(21.0)
            new_sec.page_height = Cm(29.7)
        if s.margin_header:
            new_sec.header_distance = Cm(s.margin_header)
        if s.margin_footer:
            new_sec.footer_distance = Cm(s.margin_footer)
        return new_sec

    # --- Page footer with report title + page number ---
    def setup_page_footer(self, left_text, right_text=""):
        """Set up a persistent page footer on the current (last) section.

        Args:
            left_text: Report title shown on left side
            right_text: Optional right side text. If empty, defaults to 'Page. N'
                        with automatic page numbering.
        """
        s = self.s
        section = self.doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Left side: report title
        run_left = p.add_run(left_text)
        run_left.font.name = "Arial"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = _rgb("000000")

        # Tab stop for right-alignment (add right-aligned tab)
        pPr = p._element.get_or_add_pPr()
        # Right-aligned tab at page width minus margins
        page_w = 11906 if s.paper_a4 else 12240  # A4 or Letter width in twips
        margin_l = int(s.margin_left / 2.54 * 1440)
        margin_r = int(s.margin_right / 2.54 * 1440)
        tab_pos = page_w - margin_l - margin_r
        pPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'<w:tab w:val="right" w:pos="{tab_pos}"/>'
            f'</w:tabs>'
        ))

        # Tab character
        tab_run = p.add_run("\t")
        tab_run.font.size = Pt(8)

        # Right side: "Page. " + auto page number
        run_page = p.add_run("Page. ")
        run_page.font.name = "Arial"
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = _rgb("000000")

        # PAGE field (auto page number)
        fld_xml = (
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
            f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
            f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:instrText xml:space="preserve"> PAGE \\* MERGEFORMAT </w:instrText></w:r>'
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
            f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
            f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:t>1</w:t></w:r>'
            f'<w:r {nsdecls("w")}>'
            f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
            f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
            f'<w:fldChar w:fldCharType="end"/></w:r>'
        )
        # Append field runs to the paragraph
        for r_xml in fld_xml.split('</w:r>'):
            r_xml = r_xml.strip()
            if r_xml:
                p._element.append(parse_xml(r_xml + '</w:r>'))

        # Top border on footer paragraph (separator line)
        _para_border(p, "top", 4, "000000", space=4)

        return p

    # --- Save ---
    def save(self, path):
        self.doc.save(str(path))
        return Path(path).stat().st_size


# ======================================================================
# Quick builder function
# ======================================================================
def build_word_document(content: dict, style_id: str, output_path: str) -> int:
    """Build a Word document with the specified style.

    Args:
        content: dict with 'title', 'subtitle', 'meta', 'sections'
        style_id: one of WORD_STYLES keys
        output_path: file path

    Returns:
        File size in bytes.
    """
    s = WORD_STYLES[style_id]
    w = WordBuilder(s)

    w.cover(content.get("title", ""), content.get("subtitle", ""),
            content.get("meta", ""))

    for i, sec in enumerate(content.get("sections", [])):
        w.heading(sec.get("heading", ""), level=1, number=i + 1)

        if "body" in sec:
            w.para(sec["body"])

        if "table" in sec:
            w.table(sec["table"]["headers"], sec["table"]["rows"])

        if "bullets" in sec:
            w.bullet_list(sec["bullets"])

        if "numbered" in sec:
            w.numbered_list(sec["numbered"])

        if "kv_pairs" in sec:
            w.kv_table(sec["kv_pairs"])

        if "callout" in sec:
            w.callout(sec["callout"]["text"], sec["callout"].get("kind", "NOTE"))

    w.footer(content.get("footer", ""))
    return w.save(output_path)


# ======================================================================
# Style catalog (prints summary)
# ======================================================================
def print_catalog():
    print("=" * 70)
    print("  DESIGN SYSTEM -- Available Word Styles")
    print("=" * 70)
    for sid, s in WORD_STYLES.items():
        print(f"\n  [{sid}]  {s.name}  |  {s.name_kr}")
        print(f"    {s.concept_kr}")
        print(f"    Fonts: {s.font_heading} / {s.font_body}")
        print(f"    Colors: primary={s.color_primary}, accent={s.color_accent}")
        print(f"    Cover: {s.cover_style}, Numbers: {s.number_style}, "
              f"Tables: {s.table_style}")
    print()


if __name__ == "__main__":
    print_catalog()
